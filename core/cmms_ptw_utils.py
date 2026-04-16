"""
CMMS Permit-to-Work utilities.

This module manages the Electrical Work Permit workflow for CMMS records:
- receiver application
- issuer approval
- HSE approval / permit number allocation
- closure by receiver, issuer, and HSE
- live document preview from the Word template
- downloadable filled .docx generation
"""
from __future__ import annotations

import base64
import binascii
import io
import json
import re
import uuid
from datetime import datetime
from functools import lru_cache
from pathlib import Path
from urllib.parse import parse_qs, urlencode, urlparse
from urllib.request import Request, urlopen
from zipfile import ZipFile

from django.conf import settings

BASE_DIR = Path(
    getattr(settings, 'BASE_DIR', Path(__file__).resolve().parent.parent)
).resolve()
CMMS_DATA_DIR = Path(
    getattr(settings, 'CMMS_DATA_DIR', BASE_DIR / 'cmms_data')
).resolve()
MEDIA_ROOT = Path(
    getattr(settings, 'MEDIA_ROOT', BASE_DIR / 'media')
).resolve()
MEDIA_URL = getattr(settings, 'MEDIA_URL', '/media/')

WORK_PERMIT_DIR = MEDIA_ROOT / 'work_permit'
PERMIT_TEMPLATE_PATH = WORK_PERMIT_DIR / 'ElectricalWorkPermit.docx'
PERMIT_SIGNATURES_DIR = WORK_PERMIT_DIR / 'signatures'
PERMIT_EXPORTS_DIR = WORK_PERMIT_DIR / 'generated'
PERMIT_FINAL_PDF_DIR = PERMIT_EXPORTS_DIR / 'final_pdf'
PERMIT_TEMPLATE_ASSETS_DIR = WORK_PERMIT_DIR / 'template_assets'
PERMITS_FILE = CMMS_DATA_DIR / 'permits.json'

for directory in (CMMS_DATA_DIR, WORK_PERMIT_DIR, PERMIT_SIGNATURES_DIR, PERMIT_EXPORTS_DIR, PERMIT_FINAL_PDF_DIR, PERMIT_TEMPLATE_ASSETS_DIR):
    directory.mkdir(parents=True, exist_ok=True)

PERMIT_STATUSES = {
    'draft': 'Draft',
    'pending_issue': 'Pending Issuer',
    'rejected_by_issuer': 'Rejected By Issuer',
    'pending_hse': 'Pending HSE',
    'rejected_by_hse': 'Rejected By HSE',
    'pending_receiver_number': 'Pending Receiver Permit Number',
    'active': 'Active',
    'pending_closure_issuer': 'Pending Closure Issuer',
    'closure_rejected_by_issuer': 'Closure Rejected By Issuer',
    'pending_closure_hse': 'Pending Closure HSE',
    'closure_rejected_by_hse': 'Closure Rejected By HSE',
    'closed': 'Closed',
}

LEGACY_STATUS_ALIASES = {
    'waiting_for_close': 'pending_closure_issuer',
}

SPECIAL_CELL_KEYS = {
    'receiver_acceptance': 't2_r34_c1',
    'issuer_authority': 't2_r35_c1',
    'hse_endorsement': 't2_r36_c1',
    'closure_status': 't3_r13_c1',
    'closure_receiver': 't3_r15_c1',
    'closure_issuer': 't3_r15_c5',
    'closure_hse': 't3_r17_c1',
}

SPECIAL_SIGNATURE_FIELDS = {
    SPECIAL_CELL_KEYS['receiver_acceptance']: 'receiver_signature',
    SPECIAL_CELL_KEYS['issuer_authority']: 'issuer_signature',
    SPECIAL_CELL_KEYS['hse_endorsement']: 'hse_signature',
    SPECIAL_CELL_KEYS['closure_receiver']: 'closure_receiver_signature',
    SPECIAL_CELL_KEYS['closure_issuer']: 'closure_issuer_signature',
    SPECIAL_CELL_KEYS['closure_hse']: 'closure_hse_signature',
}

APPLICATION_PREFILL_MAP = {
    't2_r2_c2': 'company_name',
    't2_r1_c9': 'start_date',
    't2_r1_c17': 'start_time',
    't2_r2_c9': 'expected_duration',
    't2_r3_c9': 'number_of_employees',
    't2_r4_c1': 'energized_equipment',
    't2_r4_c9': 'de_energized_equipment',
    't2_r5_c1': 'work_description',
    't2_r6_c1': 'location',
    't2_r7_c1': 'tools_equipment',
}

GOOGLE_PTW_TEMPLATE_URL = getattr(
    settings,
    'CMMS_PTW_TEMPLATE_URL',
    'https://docs.google.com/document/d/15aT27Vylw6SrXJq2he7EhPP9DD73D50GoIM6qSs7YJM/edit?usp=sharing',
)


def _load(path: Path):
    if not path.exists():
        return []
    try:
        with path.open('r', encoding='utf-8') as handle:
            data = json.load(handle)
            return data if isinstance(data, list) else []
    except Exception:
        return []


def _save(path: Path, data) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open('w', encoding='utf-8') as handle:
        json.dump(data, handle, indent=2, ensure_ascii=False)


def _all_permits() -> list[dict]:
    return _load(PERMITS_FILE)


def _is_deleted_permit(permit: dict | None) -> bool:
    return bool((permit or {}).get('deleted_at'))


def list_permits(*, include_deleted: bool = False) -> list[dict]:
    permits = _all_permits()
    if not include_deleted:
        permits = [permit for permit in permits if not _is_deleted_permit(permit)]
    return sorted(
        permits,
        key=lambda permit: (
            permit.get('closed_at') or '',
            permit.get('updated_at') or '',
            permit.get('created_at') or '',
            permit.get('id') or '',
        ),
        reverse=True,
    )


def get_permit(permit_id: str, *, include_deleted: bool = False) -> dict | None:
    return next((
        permit for permit in list_permits(include_deleted=include_deleted)
        if permit.get('id') == permit_id
    ), None)


def get_permit_for_record(record_id: str) -> dict | None:
    matches = [permit for permit in list_permits() if permit.get('record_id') == record_id]
    return matches[0] if matches else None


def _touch(permit: dict) -> dict:
    permit['updated_at'] = datetime.now().isoformat()
    return permit


def _save_permit(permit: dict) -> dict:
    permits = _all_permits()
    replaced = False
    for index, existing in enumerate(permits):
        if existing.get('id') == permit.get('id'):
            permits[index] = permit
            replaced = True
            break
    if not replaced:
        permits.append(permit)
    _save(PERMITS_FILE, permits)
    return permit


def update_permit(permit_id: str, data: dict) -> dict | None:
    permit = get_permit(permit_id, include_deleted=True)
    if not permit:
        return None
    permit.update({key: value for key, value in data.items() if key != 'id'})
    return _save_permit(_touch(permit))


def delete_permit(permit_id: str, deleted_by: str = '') -> dict | None:
    permit = get_permit(permit_id)
    if not permit:
        return None
    permit['deleted_at'] = datetime.now().isoformat()
    permit['deleted_by'] = str(deleted_by or '').strip()
    return _save_permit(_touch(permit))


def _slug(text: str) -> str:
    return re.sub(r'[^a-z0-9]+', '-', (text or '').lower()).strip('-') or 'permit'


def _media_url(rel_path: str | None) -> str:
    if not rel_path:
        return ''
    normalized = str(rel_path).replace('\\', '/').lstrip('/')
    return f"{MEDIA_URL.rstrip('/')}/{normalized}"


def _now_date() -> str:
    return datetime.now().strftime('%Y-%m-%d')


def _now_time() -> str:
    return datetime.now().strftime('%H:%M')


def _prefill_document_values(activity: dict | None, record: dict | None) -> dict:
    activity = activity or {}
    record = record or {}
    description = activity.get('notes') or activity.get('name') or ''
    equipment = activity.get('equipment') or activity.get('name') or ''
    technician = activity.get('assigned_technician') or ''
    prefill = {
        'company_name': 'HDEC',
        'start_date': record.get('date') or _now_date(),
        'start_time': _now_time(),
        'expected_duration': '1 Shift',
        'number_of_employees': '1' if not technician else '2',
        'energized_equipment': '',
        'de_energized_equipment': equipment,
        'work_description': description,
        'location': activity.get('location') or '',
        'tools_equipment': equipment,
    }
    return {
        cell_key: prefill.get(field_name, '')
        for cell_key, field_name in APPLICATION_PREFILL_MAP.items()
        if prefill.get(field_name, '') != ''
    }


def create_or_get_record_permit(
    record: dict,
    activity: dict,
    user: dict,
    permit_name: str = '',
    permit_link: str = '',
) -> dict:
    existing = get_permit_for_record(record.get('id', ''))
    if existing:
        return existing

    permit_id = str(uuid.uuid4())
    work_type = 'electrical'
    document_values = _prefill_document_values(activity, record)
    selected_permit_name = str(permit_name or '').strip()
    selected_permit_link = str(permit_link or '').strip()
    chosen_template_link = selected_permit_link or GOOGLE_PTW_TEMPLATE_URL
    permit = {
        'id': permit_id,
        'record_id': record.get('id', ''),
        'activity_id': record.get('activity_id', ''),
        'activity_name': record.get('activity_name', '') or activity.get('name', ''),
        'activity_description': activity.get('notes', ''),
        'activity_frequency': activity.get('frequency', ''),
        'assigned_engineer': activity.get('assigned_engineer', ''),
        'assigned_technician': activity.get('assigned_technician', ''),
        'scheduled_date': record.get('date', ''),
        'equipment': activity.get('equipment', ''),
        'location': activity.get('location', ''),
        'work_type': work_type,
        'selected_permit_name': selected_permit_name,
        'selected_permit_link': selected_permit_link,
        'template_file': str(PERMIT_TEMPLATE_PATH.relative_to(BASE_DIR)).replace('\\', '/'),
        'template_link': chosen_template_link,
        'document_link': chosen_template_link,
        'status': 'draft',
        'status_label': PERMIT_STATUSES['draft'],
        'document_values': document_values,
        'receiver_name': user.get('name', ''),
        'receiver_username': user.get('username', ''),
        'receiver_signature': '',
        'issuer_name': '',
        'issuer_username': '',
        'issuer_signature': '',
        'hse_name': '',
        'hse_username': '',
        'hse_signature': '',
        'permit_number': '',
        'receiver_confirmed_permit_number': '',
        'receiver_confirmed_at': '',
        'isolation_cert_number': '',
        'rejection_stage': '',
        'rejection_reason': '',
        'rejected_at': '',
        'rejected_by_name': '',
        'rejected_by_username': '',
        'closure_status_text': '',
        'closure_rejection_stage': '',
        'closure_rejection_reason': '',
        'closure_rejected_at': '',
        'closure_rejected_by_name': '',
        'closure_rejected_by_username': '',
        'closure_receiver_name': '',
        'closure_receiver_signature': '',
        'closure_receiver_signed_at': '',
        'closure_issuer_name': '',
        'closure_issuer_signature': '',
        'closure_issuer_signed_at': '',
        'closure_hse_name': '',
        'closure_hse_signature': '',
        'closure_hse_signed_at': '',
        'created_at': datetime.now().isoformat(),
        'updated_at': datetime.now().isoformat(),
    }
    return _save_permit(permit)


def _decode_data_url(data_url: str) -> tuple[str, bytes]:
    if not data_url or ',' not in data_url:
        raise ValueError('Invalid signature data.')
    header, encoded = data_url.split(',', 1)
    mime_match = re.match(r'data:([^;]+);base64$', header)
    if not mime_match:
        raise ValueError('Unsupported signature format.')
    mime_type = mime_match.group(1).lower()
    ext_map = {
        'image/png': '.png',
        'image/jpeg': '.jpg',
        'image/jpg': '.jpg',
        'image/webp': '.webp',
    }
    ext = ext_map.get(mime_type, '.png')
    try:
        payload = base64.b64decode(encoded)
    except (ValueError, binascii.Error) as exc:
        raise ValueError('Invalid signature data.') from exc
    return ext, payload


def save_signature_image(permit_id: str, field_name: str, data_url: str | None) -> str:
    if not data_url:
        return ''
    ext, payload = _decode_data_url(data_url)
    dest_dir = PERMIT_SIGNATURES_DIR / permit_id
    dest_dir.mkdir(parents=True, exist_ok=True)
    dest = dest_dir / f'{field_name}{ext}'
    dest.write_bytes(payload)
    return str(dest.relative_to(MEDIA_ROOT)).replace('\\', '/')


def can_edit_application(permit: dict, user: dict | None) -> bool:
    role = (user or {}).get('role', '')
    if role == 'admin':
        return permit.get('status') in ('draft', 'pending_issue', 'rejected_by_issuer', 'rejected_by_hse')
    return (
        role == 'maintenance_engineer'
        and permit.get('status') in ('draft', 'pending_issue', 'rejected_by_issuer', 'rejected_by_hse')
        and permit.get('receiver_username') == (user or {}).get('username', '')
    )


def can_issue_permit(permit: dict, user: dict | None) -> bool:
    role = (user or {}).get('role', '')
    return permit.get('status') == 'pending_issue' and role in ('admin', 'operation_engineer')


def can_hse_approve(permit: dict, user: dict | None) -> bool:
    role = (user or {}).get('role', '')
    return permit.get('status') == 'pending_hse' and role in ('admin', 'hse_engineer')


def can_receiver_unlock(permit: dict, user: dict | None) -> bool:
    role = (user or {}).get('role', '')
    if role == 'admin':
        return permit.get('status') == 'pending_receiver_number'
    return (
        role == 'maintenance_engineer'
        and permit.get('status') == 'pending_receiver_number'
        and permit.get('receiver_username') == (user or {}).get('username', '')
    )


def can_close_receiver(permit: dict, user: dict | None) -> bool:
    role = (user or {}).get('role', '')
    if role == 'admin':
        return permit.get('status') in ('active', 'closure_rejected_by_issuer', 'closure_rejected_by_hse')
    return (
        role == 'maintenance_engineer'
        and permit.get('status') in ('active', 'closure_rejected_by_issuer', 'closure_rejected_by_hse')
        and permit.get('receiver_username') == (user or {}).get('username', '')
    )


def can_close_issuer(permit: dict, user: dict | None) -> bool:
    role = (user or {}).get('role', '')
    return permit.get('status') == 'pending_closure_issuer' and role in ('admin', 'operation_engineer')


def can_close_hse(permit: dict, user: dict | None) -> bool:
    role = (user or {}).get('role', '')
    return permit.get('status') == 'pending_closure_hse' and role in ('admin', 'hse_engineer')


def can_delete_permit(permit: dict, user: dict | None) -> bool:
    if not permit or not is_cmms_permit(permit):
        return False
    role = (user or {}).get('role', '')
    username = str((user or {}).get('username', '')).strip()
    if role == 'admin':
        return True
    return (
        role == 'maintenance_engineer'
        and username
        and permit.get('receiver_username') == username
        and permit.get('status') not in (
            'active',
            'pending_closure_issuer',
            'pending_closure_hse',
            'closure_rejected_by_issuer',
            'closure_rejected_by_hse',
        )
    )


def application_is_active(permit: dict) -> bool:
    return permit.get('status') in (
        'active',
        'pending_closure_issuer',
        'pending_closure_hse',
        'closure_rejected_by_issuer',
        'closure_rejected_by_hse',
        'closed',
    )


def is_cmms_permit(permit: dict | None) -> bool:
    return bool((permit or {}).get('record_id'))


def _twips_to_px(value) -> int:
    try:
        return max(30, int(round(int(value) / 15)))
    except Exception:
        return 90


def _emu_to_px(value) -> int:
    try:
        return max(1, int(round(int(value) / 9525)))
    except Exception:
        return 24


def _get_cell_fill(cell) -> str:
    try:
        shd = cell._tc.tcPr.shd
        fill = shd.fill if shd is not None else ''
        if fill and fill.lower() not in ('auto', 'ffffff', '000000', 'transparent'):
            return f'#{fill}'
    except Exception:
        pass
    return ''


def _paragraph_alignment(paragraph) -> str:
    try:
        alignment = paragraph.alignment
    except Exception:
        alignment = None
    mapping = {
        0: 'left',
        1: 'center',
        2: 'right',
        3: 'justify',
    }
    return mapping.get(getattr(alignment, 'value', None), 'left')


def _paragraph_spacing(paragraph) -> dict:
    try:
        fmt = paragraph.paragraph_format
    except Exception:
        return {'before': 0, 'after': 0}
    before = 0
    after = 0
    try:
        if fmt.space_before:
            before = max(0, int(round(float(fmt.space_before.pt) * 1.333)))
    except Exception:
        before = 0
    try:
        if fmt.space_after:
            after = max(0, int(round(float(fmt.space_after.pt) * 1.333)))
    except Exception:
        after = 0
    return {'before': before, 'after': after}


def _get_cell_alignment(cell) -> str:
    try:
        alignment = cell.paragraphs[0].alignment
    except Exception:
        alignment = None
    mapping = {
        0: 'left',
        1: 'center',
        2: 'right',
        3: 'justify',
    }
    return mapping.get(getattr(alignment, 'value', None), 'left')


def _get_cell_vertical(cell) -> str:
    try:
        vertical = cell.vertical_alignment
    except Exception:
        vertical = None
    mapping = {
        0: 'top',
        1: 'center',
        3: 'bottom',
    }
    return mapping.get(getattr(vertical, 'value', None), 'top')


def _first_font(cell) -> dict:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            text = (run.text or '').strip()
            if not text:
                continue
            return {
                'bold': bool(run.bold),
                'italic': bool(run.italic),
                'size': int(round(float(run.font.size.pt))) if run.font.size else None,
            }
    return {'bold': False, 'italic': False, 'size': None}


def _extract_template_assets(document) -> dict[str, str]:
    if not PERMIT_TEMPLATE_PATH.exists():
        return {}
    asset_dir = PERMIT_TEMPLATE_ASSETS_DIR / PERMIT_TEMPLATE_PATH.stem
    asset_dir.mkdir(parents=True, exist_ok=True)
    assets: dict[str, str] = {}
    with ZipFile(PERMIT_TEMPLATE_PATH) as archive:
        for name in archive.namelist():
            if not name.startswith('word/media/') or name.endswith('/'):
                continue
            filename = Path(name).name
            dest = asset_dir / filename
            if not dest.exists():
                dest.write_bytes(archive.read(name))
            rel = document.part.rels
            _ = rel
            assets[f"media/{filename}"] = _media_url(str(dest.relative_to(MEDIA_ROOT)).replace('\\', '/'))
    return assets


def _run_style(run) -> dict:
    color = ''
    try:
        if run.font.color and run.font.color.rgb:
            color = f"#{run.font.color.rgb}"
    except Exception:
        color = ''
    font_name = ''
    try:
        font_name = run.font.name or ''
    except Exception:
        font_name = ''
    font_size = None
    try:
        if run.font.size:
            font_size = int(round(float(run.font.size.pt)))
    except Exception:
        font_size = None
    return {
        'bold': bool(run.bold),
        'italic': bool(run.italic),
        'underline': bool(run.underline),
        'font_size': font_size,
        'font_name': font_name,
        'color': color,
    }


def _run_items(run, image_map: dict[str, str]) -> list[dict]:
    from docx.oxml.ns import qn

    items: list[dict] = []
    style = _run_style(run)
    for child in run._element:
        if child.tag == qn('w:t'):
            items.append({
                'type': 'text',
                'text': child.text or '',
                **style,
            })
        elif child.tag == qn('w:tab'):
            items.append({
                'type': 'text',
                'text': '    ',
                **style,
            })
        elif child.tag in (qn('w:br'), qn('w:cr')):
            items.append({
                'type': 'text',
                'text': '\n',
                **style,
            })
        elif child.tag == qn('w:drawing'):
            blips = list(child.iter(qn('a:blip')))
            for blip in blips:
                rel_id = blip.get(qn('r:embed'))
                if not rel_id:
                    continue
                rel = run.part.rels.get(rel_id)
                if not rel:
                    continue
                src = image_map.get(rel.target_ref, '')
                if not src:
                    continue
                extent = list(child.iter(qn('wp:extent')))
                width = _emu_to_px(extent[0].get('cx')) if extent else 64
                height = _emu_to_px(extent[0].get('cy')) if extent else 64
                items.append({
                    'type': 'image',
                    'src': src,
                    'width': width,
                    'height': height,
                })
    return items


def _serialize_paragraph(paragraph, image_map: dict[str, str], *, keep_empty: bool = False) -> dict | None:
    items: list[dict] = []
    for run in paragraph.runs:
        items.extend(_run_items(run, image_map))
    has_content = any(
        item.get('type') == 'image' or str(item.get('text', '')).strip()
        for item in items
    )
    if not has_content and not keep_empty:
        return None
    spacing = _paragraph_spacing(paragraph)
    return {
        'items': items,
        'align': _paragraph_alignment(paragraph),
        'space_before': spacing['before'],
        'space_after': spacing['after'],
        'empty': not has_content,
    }


def _serialize_cell_paragraphs(cell, image_map: dict[str, str]) -> list[dict]:
    paragraphs: list[dict] = []
    for paragraph in cell.paragraphs:
        data = _serialize_paragraph(paragraph, image_map, keep_empty=True)
        if data is not None:
            paragraphs.append(data)
    return paragraphs


def _paragraphs_have_visible_content(paragraphs: list[dict]) -> bool:
    return any(
        item.get('type') == 'image' or str(item.get('text', '')).strip()
        for paragraph in paragraphs
        for item in paragraph.get('items', [])
    )


def _cell_borders(cell) -> dict:
    borders = {}
    try:
        tc_borders = cell._tc.tcPr.tcBorders
    except Exception:
        tc_borders = None
    if tc_borders is None:
        return borders
    for side in ('top', 'left', 'bottom', 'right'):
        border = getattr(tc_borders, side, None)
        if border is None:
            continue
        val = border.val or ''
        if val in ('nil', 'none'):
            borders[side] = 'none'
            continue
        color = border.color if border.color and border.color.lower() != 'auto' else '000000'
        try:
            width = max(1, int(round(int(border.sz or 8) / 8)))
        except Exception:
            width = 1
        borders[side] = f'{width}px solid #{color}'
    return borders


def _serialize_table(table_index: int, table, image_map: dict[str, str]) -> dict:
    col_count = len(table.columns)
    matrix = [[row.cells[col] for col in range(col_count)] for row in table.rows]
    grid = []
    tbl_grid = getattr(table._tbl, 'tblGrid', None)
    if tbl_grid is not None:
        grid = [_twips_to_px(col.w) for col in tbl_grid.gridCol_lst]
    if len(grid) < col_count:
        grid.extend([90] * (col_count - len(grid)))

    row_heights = []
    for row in table.rows:
        try:
            height = _twips_to_px(row.height.twips) if row.height else 0
        except Exception:
            height = 0
        row_heights.append(height)

    seen: set[tuple[int, int]] = set()
    cells: list[dict] = []
    for row_index in range(len(matrix)):
        for col_index in range(col_count):
            if (row_index, col_index) in seen:
                continue
            cell = matrix[row_index][col_index]
            colspan = 1
            while (
                col_index + colspan < col_count
                and matrix[row_index][col_index + colspan]._tc is cell._tc
            ):
                colspan += 1
            rowspan = 1
            while row_index + rowspan < len(matrix):
                if all(
                    matrix[row_index + rowspan][check_col]._tc is cell._tc
                    for check_col in range(col_index, col_index + colspan)
                ):
                    rowspan += 1
                else:
                    break
            for r in range(row_index, row_index + rowspan):
                for c in range(col_index, col_index + colspan):
                    seen.add((r, c))

            paragraphs = _serialize_cell_paragraphs(cell, image_map)
            font = _first_font(cell)
            cell_key = f't{table_index}_r{row_index + 1}_c{col_index + 1}'
            cells.append({
                'key': cell_key,
                'row': row_index + 1,
                'col': col_index + 1,
                'rowspan': rowspan,
                'colspan': colspan,
                'paragraphs': paragraphs,
                'blank': not _paragraphs_have_visible_content(paragraphs),
                'background': _get_cell_fill(cell),
                'text_align': _get_cell_alignment(cell),
                'vertical_align': _get_cell_vertical(cell),
                'bold': font['bold'],
                'italic': font['italic'],
                'font_size': font['size'],
                'borders': _cell_borders(cell),
                'min_height': row_heights[row_index] if row_index < len(row_heights) else 0,
            })

    return {
        'type': 'table',
        'index': table_index,
        'column_widths': grid,
        'row_heights': row_heights,
        'cells': cells,
    }


@lru_cache(maxsize=1)
def load_template_schema() -> dict:
    from docx import Document
    from docx.oxml.ns import qn

    if not PERMIT_TEMPLATE_PATH.exists():
        return {'blocks': []}

    document = Document(str(PERMIT_TEMPLATE_PATH))
    image_map = _extract_template_assets(document)
    blocks: list[dict] = []
    paragraph_index = 0
    table_index = 0
    for child in document.element.body.iterchildren():
        if child.tag == qn('w:p'):
            paragraph = document.paragraphs[paragraph_index]
            paragraph_index += 1
            blocks.append({
                'type': 'paragraph',
                **(_serialize_paragraph(paragraph, image_map, keep_empty=True) or {
                    'items': [],
                    'align': 'left',
                    'space_before': 0,
                    'space_after': 0,
                    'empty': True,
                }),
            })
        elif child.tag == qn('w:tbl'):
            table = document.tables[table_index]
            table_index += 1
            blocks.append(_serialize_table(table_index, table, image_map))
    return {'blocks': blocks}


def _special_document_blocks(
    permit: dict,
    *,
    receiver_signable: bool = False,
    issuer_signable: bool = False,
    hse_signable: bool = False,
    closure_text_editable: bool = False,
    closure_receiver_signable: bool = False,
    closure_issuer_signable: bool = False,
    closure_hse_signable: bool = False,
) -> dict:
    special = {}

    receiver_lines = []
    if permit.get('receiver_name'):
        receiver_lines.append(f"Permit Receiver Name: {permit.get('receiver_name')}")
    if permit.get('submitted_at'):
        receiver_lines.append(f"Date: {permit.get('submitted_at')[:16].replace('T', ' ')}")
    if receiver_lines or permit.get('receiver_signature') or receiver_signable:
        special[SPECIAL_CELL_KEYS['receiver_acceptance']] = {
            'lines': receiver_lines,
            'image_url': _media_url(permit.get('receiver_signature')),
            'signature_field': 'receiver_signature',
            'signature_label': 'Receiver Signature',
            'signature_editable': receiver_signable,
        }

    issuer_lines = []
    if permit.get('issuer_name'):
        issuer_lines.append(f"Permit Issuer Name: {permit.get('issuer_name')}")
    if permit.get('issued_at'):
        issuer_lines.append(f"Date: {permit.get('issued_at')[:16].replace('T', ' ')}")
    if issuer_lines or permit.get('issuer_signature') or issuer_signable:
        special[SPECIAL_CELL_KEYS['issuer_authority']] = {
            'lines': issuer_lines,
            'image_url': _media_url(permit.get('issuer_signature')),
            'signature_field': 'issuer_signature',
            'signature_label': 'Issuer Signature',
            'signature_editable': issuer_signable,
        }

    hse_lines = []
    if permit.get('hse_name'):
        hse_lines.append(f"Name: {permit.get('hse_name')}")
    if permit.get('permit_number'):
        hse_lines.append(f"PTW Ref. No: {permit.get('permit_number')}")
    if permit.get('hse_signed_at'):
        hse_lines.append(f"Date: {permit.get('hse_signed_at')[:16].replace('T', ' ')}")
    if hse_lines or permit.get('hse_signature') or hse_signable:
        special[SPECIAL_CELL_KEYS['hse_endorsement']] = {
            'lines': hse_lines,
            'image_url': _media_url(permit.get('hse_signature')),
            'signature_field': 'hse_signature',
            'signature_label': 'HSE Signature',
            'signature_editable': hse_signable,
        }

    closure_status_lines = []
    if permit.get('closure_requested_at'):
        closure_status_lines.append(f"Submitted: {permit.get('closure_requested_at')[:16].replace('T', ' ')}")
    if closure_status_lines or permit.get('closure_status_text') or closure_text_editable:
        special[SPECIAL_CELL_KEYS['closure_status']] = {
            'lines': closure_status_lines,
            'text_value': permit.get('closure_status_text', ''),
            'text_editable': closure_text_editable,
            'text_placeholder': 'Enter closure details directly on the permit',
        }

    closure_receiver_lines = []
    if permit.get('closure_receiver_name'):
        closure_receiver_lines.append(f"Name: {permit.get('closure_receiver_name')}")
    if permit.get('closure_receiver_signed_at'):
        closure_receiver_lines.append(f"Date: {permit.get('closure_receiver_signed_at')[:16].replace('T', ' ')}")
    if closure_receiver_lines or permit.get('closure_receiver_signature') or closure_receiver_signable:
        special[SPECIAL_CELL_KEYS['closure_receiver']] = {
            'lines': closure_receiver_lines,
            'image_url': _media_url(permit.get('closure_receiver_signature')),
            'signature_field': 'closure_receiver_signature',
            'signature_label': 'Receiver Closure Signature',
            'signature_editable': closure_receiver_signable,
        }

    closure_issuer_lines = []
    if permit.get('closure_issuer_name'):
        closure_issuer_lines.append(f"Name: {permit.get('closure_issuer_name')}")
    if permit.get('closure_issuer_signed_at'):
        closure_issuer_lines.append(f"Date: {permit.get('closure_issuer_signed_at')[:16].replace('T', ' ')}")
    if closure_issuer_lines or permit.get('closure_issuer_signature') or closure_issuer_signable:
        special[SPECIAL_CELL_KEYS['closure_issuer']] = {
            'lines': closure_issuer_lines,
            'image_url': _media_url(permit.get('closure_issuer_signature')),
            'signature_field': 'closure_issuer_signature',
            'signature_label': 'Issuer Closure Signature',
            'signature_editable': closure_issuer_signable,
        }

    closure_hse_lines = []
    if permit.get('closure_hse_name'):
        closure_hse_lines.append(f"Name: {permit.get('closure_hse_name')}")
    if permit.get('closure_hse_signed_at'):
        closure_hse_lines.append(f"Date: {permit.get('closure_hse_signed_at')[:16].replace('T', ' ')}")
    if closure_hse_lines or permit.get('closure_hse_signature') or closure_hse_signable:
        special[SPECIAL_CELL_KEYS['closure_hse']] = {
            'lines': closure_hse_lines,
            'image_url': _media_url(permit.get('closure_hse_signature')),
            'signature_field': 'closure_hse_signature',
            'signature_label': 'HSE Closure Signature',
            'signature_editable': closure_hse_signable,
        }
    return special


def build_document_payload(
    permit: dict,
    *,
    application_editable: bool = False,
    issuer_signable: bool = False,
    hse_editable: bool = False,
    closure_text_editable: bool = False,
    closure_issuer_signable: bool = False,
    closure_hse_signable: bool = False,
) -> dict:
    schema = load_template_schema()
    doc_values = permit.get('document_values', {})
    special = _special_document_blocks(
        permit,
        receiver_signable=application_editable,
        issuer_signable=issuer_signable,
        hse_signable=hse_editable,
        closure_text_editable=closure_text_editable,
        closure_receiver_signable=closure_text_editable,
        closure_issuer_signable=closure_issuer_signable,
        closure_hse_signable=closure_hse_signable,
    )
    explicit_editable_keys = set()
    if hse_editable:
        explicit_editable_keys.update({'t2_r1_c2', 't2_r3_c2'})

    blocks = []
    for block in schema.get('blocks', []):
        if block.get('type') != 'table':
            blocks.append(block)
            continue
        rendered_cells = []
        for cell in block.get('cells', []):
            key = cell['key']
            rendered_cells.append({
                **cell,
                'value': doc_values.get(key, ''),
                'editable': bool((application_editable and cell.get('blank')) or key in explicit_editable_keys),
                'special': special.get(key, {}),
            })
        blocks.append({
            **block,
            'cells': rendered_cells,
        })
    return {'blocks': blocks}


def _clear_cell(cell) -> None:
    tc = cell._tc
    for child in list(tc):
        if child.tag.endswith('tcPr'):
            continue
        tc.remove(child)
    cell.add_paragraph('')


def _set_cell_text(cell, value: str) -> None:
    _clear_cell(cell)
    paragraph = cell.paragraphs[0]
    paragraph.text = value or ''


def _append_lines(cell, lines: list[str], image_path: Path | None = None) -> None:
    if lines:
        for line in lines:
            paragraph = cell.add_paragraph()
            paragraph.text = line
    if image_path and image_path.exists():
        paragraph = cell.add_paragraph()
        run = paragraph.add_run()
        try:
            from docx.shared import Inches

            run.add_picture(str(image_path), width=Inches(1.2))
        except Exception:
            paragraph.add_run(f'[Signature: {image_path.name}]')


def _apply_special_sections(document, permit: dict) -> None:
    table2 = document.tables[1] if len(document.tables) > 1 else None
    table3 = document.tables[2] if len(document.tables) > 2 else None
    if table2:
        if permit.get('receiver_name') or permit.get('submitted_at') or permit.get('receiver_signature'):
            cell = table2.rows[33].cells[0]
            _append_lines(
                cell,
                [line for line in [
                    f"Permit Receiver Name: {permit.get('receiver_name')}" if permit.get('receiver_name') else '',
                    f"Date: {permit.get('submitted_at')[:16].replace('T', ' ')}" if permit.get('submitted_at') else '',
                ] if line],
                (MEDIA_ROOT / permit['receiver_signature']) if permit.get('receiver_signature') else None,
            )
        if permit.get('issuer_name') or permit.get('issued_at') or permit.get('issuer_signature'):
            cell = table2.rows[34].cells[0]
            _append_lines(
                cell,
                [line for line in [
                    f"Permit Issuer Name: {permit.get('issuer_name')}" if permit.get('issuer_name') else '',
                    f"Date: {permit.get('issued_at')[:16].replace('T', ' ')}" if permit.get('issued_at') else '',
                ] if line],
                (MEDIA_ROOT / permit['issuer_signature']) if permit.get('issuer_signature') else None,
            )
        if permit.get('hse_name') or permit.get('hse_signed_at') or permit.get('hse_signature') or permit.get('permit_number'):
            cell = table2.rows[35].cells[0]
            _append_lines(
                cell,
                [line for line in [
                    f"Name: {permit.get('hse_name')}" if permit.get('hse_name') else '',
                    f"PTW Ref. No: {permit.get('permit_number')}" if permit.get('permit_number') else '',
                    f"Date: {permit.get('hse_signed_at')[:16].replace('T', ' ')}" if permit.get('hse_signed_at') else '',
                ] if line],
                (MEDIA_ROOT / permit['hse_signature']) if permit.get('hse_signature') else None,
            )
    if table3:
        if permit.get('closure_status_text') or permit.get('closure_requested_at'):
            cell = table3.rows[12].cells[0]
            _append_lines(
                cell,
                [line for line in [
                    permit.get('closure_status_text', ''),
                    f"Submitted: {permit.get('closure_requested_at')[:16].replace('T', ' ')}" if permit.get('closure_requested_at') else '',
                ] if line],
            )
        if permit.get('closure_receiver_name') or permit.get('closure_receiver_signed_at') or permit.get('closure_receiver_signature'):
            cell = table3.rows[14].cells[0]
            _append_lines(
                cell,
                [line for line in [
                    f"Name: {permit.get('closure_receiver_name')}" if permit.get('closure_receiver_name') else '',
                    f"Date: {permit.get('closure_receiver_signed_at')[:16].replace('T', ' ')}" if permit.get('closure_receiver_signed_at') else '',
                ] if line],
                (MEDIA_ROOT / permit['closure_receiver_signature']) if permit.get('closure_receiver_signature') else None,
            )
        if permit.get('closure_issuer_name') or permit.get('closure_issuer_signed_at') or permit.get('closure_issuer_signature'):
            cell = table3.rows[14].cells[4]
            _append_lines(
                cell,
                [line for line in [
                    f"Name: {permit.get('closure_issuer_name')}" if permit.get('closure_issuer_name') else '',
                    f"Date: {permit.get('closure_issuer_signed_at')[:16].replace('T', ' ')}" if permit.get('closure_issuer_signed_at') else '',
                ] if line],
                (MEDIA_ROOT / permit['closure_issuer_signature']) if permit.get('closure_issuer_signature') else None,
            )
        if permit.get('closure_hse_name') or permit.get('closure_hse_signed_at') or permit.get('closure_hse_signature'):
            cell = table3.rows[16].cells[0]
            _append_lines(
                cell,
                [line for line in [
                    f"Name: {permit.get('closure_hse_name')}" if permit.get('closure_hse_name') else '',
                    f"Date: {permit.get('closure_hse_signed_at')[:16].replace('T', ' ')}" if permit.get('closure_hse_signed_at') else '',
                ] if line],
                (MEDIA_ROOT / permit['closure_hse_signature']) if permit.get('closure_hse_signature') else None,
            )


def _download_google_export(url: str, export_format: str = 'pdf') -> bytes | None:
    link = str(url or '').strip()
    if not link:
        return None
    try:
        parsed = urlparse(link)
    except Exception:
        return None
    if parsed.hostname != 'docs.google.com':
        return None

    document_match = re.search(r'/document/d/([a-zA-Z0-9\-_]+)', parsed.path or '')
    if document_match:
        export_url = f"https://docs.google.com/document/d/{document_match.group(1)}/export?{urlencode({'format': export_format})}"
    else:
        spreadsheet_match = re.search(r'/spreadsheets/d/([a-zA-Z0-9\-_]+)', parsed.path or '')
        if not spreadsheet_match:
            return None
        query = parse_qs(parsed.query or '')
        fragment = parse_qs(str(parsed.fragment or '').lstrip('#'))
        gid = str(query.get('gid', [''])[0] or fragment.get('gid', [''])[0]).strip()
        params = {'format': export_format}
        if export_format == 'pdf':
            params.update({
                'portrait': 'false',
                'size': 'A4',
                'fitw': 'true',
                'sheetnames': 'false',
                'printtitle': 'false',
                'pagenumbers': 'false',
                'gridlines': 'false',
                'fzr': 'false',
            })
        if gid:
            params['gid'] = gid
        export_url = f"https://docs.google.com/spreadsheets/d/{spreadsheet_match.group(1)}/export?{urlencode(params)}"

    try:
        request = Request(export_url, headers={'User-Agent': 'HDEC-CMMS/1.0'})
        with urlopen(request, timeout=15) as response:
            payload = response.read()
        if not payload:
            return None
        if payload[:256].lower().startswith(b'<!doctype html') or payload[:256].lower().startswith(b'<html'):
            return None
        return payload
    except Exception:
        return None


def build_permit_pdf(permit: dict) -> io.BytesIO:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.utils import simpleSplit
    from reportlab.pdfgen import canvas

    page_width, page_height = A4
    margin = 18 * mm
    line_height = 13
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    y = page_height - margin

    def header():
        nonlocal y
        pdf.setFillColorRGB(0.09, 0.25, 0.54)
        pdf.rect(0, page_height - 36 * mm, page_width, 36 * mm, stroke=0, fill=1)
        pdf.setFillColorRGB(1, 1, 1)
        pdf.setFont('Helvetica-Bold', 18)
        pdf.drawString(margin, page_height - 18 * mm, 'HDEC CMMS Permit To Work')
        pdf.setFont('Helvetica', 10)
        pdf.drawString(margin, page_height - 25 * mm, f"Activity: {permit.get('activity_name') or permit.get('equipment') or 'Permit'}")
        pdf.drawRightString(page_width - margin, page_height - 25 * mm, f"Permit No: {permit.get('permit_number') or 'N/A'}")
        y = page_height - 45 * mm

    def new_page():
        pdf.showPage()
        header()

    def ensure_space(lines_needed: int = 2):
        nonlocal y
        if y - (lines_needed * line_height) < margin:
            new_page()

    def section(title: str):
        nonlocal y
        ensure_space(3)
        pdf.setFillColorRGB(0.11, 0.17, 0.27)
        pdf.setFont('Helvetica-Bold', 12)
        pdf.drawString(margin, y, title)
        y -= 6
        pdf.setStrokeColorRGB(0.86, 0.89, 0.94)
        pdf.line(margin, y, page_width - margin, y)
        y -= 14

    def field(label: str, value: str):
        nonlocal y
        text = str(value or 'N/A')
        wrapped = simpleSplit(text, 'Helvetica', 10, page_width - (2 * margin) - 100)
        ensure_space(max(2, len(wrapped) + 1))
        pdf.setFillColorRGB(0.20, 0.26, 0.35)
        pdf.setFont('Helvetica-Bold', 10)
        pdf.drawString(margin, y, label)
        pdf.setFillColorRGB(0.07, 0.09, 0.12)
        pdf.setFont('Helvetica', 10)
        current_y = y
        for index, line in enumerate(wrapped or ['N/A']):
            pdf.drawString(margin + 100, current_y - (index * line_height), line)
        y = current_y - (max(1, len(wrapped)) * line_height) - 4

    header()
    section('Permit Summary')
    for label, value in [
        ('Status', permit.get('status_label') or permit.get('status') or 'N/A'),
        ('Receiver', permit.get('receiver_name') or 'N/A'),
        ('Issuer', permit.get('issuer_name') or 'N/A'),
        ('HSE', permit.get('closure_hse_name') or permit.get('hse_name') or 'N/A'),
        ('Equipment', permit.get('equipment') or 'N/A'),
        ('Location', permit.get('location') or 'N/A'),
        ('Scheduled Date', permit.get('scheduled_date') or 'N/A'),
        ('Work Type', (permit.get('work_type') or '').replace('_', ' ').title() or 'N/A'),
        ('Isolation Cert.', permit.get('isolation_cert_number') or 'N/A'),
        ('Google Permit Link', permit.get('document_link') or 'N/A'),
    ]:
        field(label, value)

    section('Workflow Sign-Off')
    for label, value in [
        ('Receiver Submitted', permit.get('submitted_at') or 'N/A'),
        ('Issuer Approved', permit.get('issued_at') or 'N/A'),
        ('HSE Approved', permit.get('hse_signed_at') or 'N/A'),
        ('Receiver Proceeded', permit.get('receiver_confirmed_at') or 'N/A'),
        ('Closure Requested', permit.get('closure_requested_at') or 'N/A'),
        ('Issuer Closure', permit.get('closure_issuer_signed_at') or 'N/A'),
        ('HSE Closure', permit.get('closure_hse_signed_at') or 'N/A'),
        ('Closed At', permit.get('closed_at') or 'N/A'),
    ]:
        field(label, value)

    closure_notes = str(permit.get('closure_status_text') or '').strip()
    if closure_notes:
        section('Closure Notes')
        field('Notes', closure_notes)

    pdf.setFont('Helvetica-Oblique', 9)
    pdf.setFillColorRGB(0.38, 0.45, 0.54)
    ensure_space(2)
    pdf.drawString(margin, y, 'This PDF is the final locked CMMS permit export generated after closure.')
    pdf.save()
    buffer.seek(0)
    return buffer


def ensure_final_permit_pdf(permit: dict | None) -> dict | None:
    annotated = annotate_permit(permit)
    if not annotated:
        return None

    pdf_rel = str(annotated.get('final_pdf', '') or '').strip()
    if pdf_rel:
        pdf_path = MEDIA_ROOT / pdf_rel
        if pdf_path.exists():
            return annotated

    pdf_bytes = _download_google_export(annotated.get('document_link') or annotated.get('template_link'), 'pdf')
    if not pdf_bytes:
        pdf_bytes = build_permit_pdf(annotated).getvalue()
    if not pdf_bytes:
        return annotated

    filename = permit_filename(annotated).replace('.docx', '.pdf')
    dest_dir = PERMIT_FINAL_PDF_DIR / str(annotated.get('id', 'permit'))
    dest_dir.mkdir(parents=True, exist_ok=True)
    dest_path = dest_dir / filename
    dest_path.write_bytes(pdf_bytes)
    rel_path = str(dest_path.relative_to(MEDIA_ROOT)).replace('\\', '/')
    updated = update_permit(annotated['id'], {
        'final_pdf': rel_path,
        'finalized_at': annotated.get('closed_at') or datetime.now().isoformat(),
    })
    return annotate_permit(updated)


def build_permit_docx(permit: dict) -> io.BytesIO:
    from docx import Document

    if not PERMIT_TEMPLATE_PATH.exists():
        raise FileNotFoundError('ElectricalWorkPermit.docx not found.')

    document = Document(str(PERMIT_TEMPLATE_PATH))
    for cell_key, value in (permit.get('document_values') or {}).items():
        match = re.match(r'^t(\d+)_r(\d+)_c(\d+)$', str(cell_key))
        if not match:
            continue
        table_index, row_index, col_index = (int(part) for part in match.groups())
        try:
            table = document.tables[table_index - 1]
            cell = table.rows[row_index - 1].cells[col_index - 1]
        except Exception:
            continue
        _set_cell_text(cell, str(value or ''))

    _apply_special_sections(document, permit)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def permit_filename(permit: dict) -> str:
    number = permit.get('permit_number') or permit.get('id') or 'permit'
    date_part = permit.get('scheduled_date') or _now_date()
    slug = _slug(permit.get('activity_name') or permit.get('equipment') or 'electrical-permit')
    return f'{slug}_{number}_{date_part}.docx'.replace(' ', '_')


def annotate_permit(permit: dict | None) -> dict | None:
    if not permit:
        return None
    enriched = dict(permit)
    status = str(permit.get('status', '') or '')
    normalized_status = LEGACY_STATUS_ALIASES.get(status, status)
    enriched['status'] = normalized_status
    enriched['status_label'] = PERMIT_STATUSES.get(
        normalized_status,
        normalized_status.replace('_', ' ').title() if normalized_status else '',
    )
    enriched['receiver_signature_url'] = _media_url(permit.get('receiver_signature'))
    enriched['issuer_signature_url'] = _media_url(permit.get('issuer_signature'))
    enriched['hse_signature_url'] = _media_url(permit.get('hse_signature'))
    enriched['closure_receiver_signature_url'] = _media_url(permit.get('closure_receiver_signature'))
    enriched['closure_issuer_signature_url'] = _media_url(permit.get('closure_issuer_signature'))
    enriched['closure_hse_signature_url'] = _media_url(permit.get('closure_hse_signature'))
    enriched['final_pdf_url'] = _media_url(permit.get('final_pdf'))
    enriched['template_link'] = permit.get('template_link') or GOOGLE_PTW_TEMPLATE_URL
    enriched['document_link'] = permit.get('document_link') or enriched['template_link']
    return enriched
